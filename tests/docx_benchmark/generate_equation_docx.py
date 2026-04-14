#!/usr/bin/env python3
"""Generate a DOCX file with OMML equations to validate MDM's math support.

OMML (Office Math Markup Language) is embedded via direct XML injection
because python-docx has no native math API.
"""

import os
from docx import Document
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement, parse_xml

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

M = "http://schemas.openxmlformats.org/officeDocument/2006/math"

OMML_EINSTEIN = f"""<m:oMathPara xmlns:m="{M}">
  <m:oMath>
    <m:sSup>
      <m:e><m:r><m:t>E=mc</m:t></m:r></m:e>
      <m:sup><m:r><m:t>2</m:t></m:r></m:sup>
    </m:sSup>
  </m:oMath>
</m:oMathPara>"""

OMML_INLINE_FRAC = f"""<m:oMath xmlns:m="{M}">
  <m:f>
    <m:num><m:r><m:t>1</m:t></m:r></m:num>
    <m:den><m:r><m:t>2</m:t></m:r></m:den>
  </m:f>
</m:oMath>"""

OMML_SUM = f"""<m:oMathPara xmlns:m="{M}">
  <m:oMath>
    <m:nary>
      <m:naryPr><m:chr m:val="\u2211"/></m:naryPr>
      <m:sub><m:r><m:t>i=1</m:t></m:r></m:sub>
      <m:sup><m:r><m:t>n</m:t></m:r></m:sup>
      <m:e>
        <m:sSup>
          <m:e><m:r><m:t>i</m:t></m:r></m:e>
          <m:sup><m:r><m:t>2</m:t></m:r></m:sup>
        </m:sSup>
      </m:e>
    </m:nary>
  </m:oMath>
</m:oMathPara>"""

OMML_SQRT = f"""<m:oMath xmlns:m="{M}">
  <m:rad>
    <m:deg><m:r><m:t>3</m:t></m:r></m:deg>
    <m:e><m:r><m:t>x</m:t></m:r></m:e>
  </m:rad>
</m:oMath>"""


def inject_math_block(doc, omml_xml):
    """Append a paragraph whose body is a block-level oMathPara."""
    p = doc.add_paragraph()
    p._p.append(parse_xml(omml_xml))
    return p


def inject_math_inline(doc, prefix_text, omml_xml, suffix_text):
    """Add a paragraph with text around an inline oMath."""
    p = doc.add_paragraph(prefix_text)
    p._p.append(parse_xml(omml_xml))
    if suffix_text:
        p.add_run(suffix_text)
    return p


def create_equation_test():
    doc = Document()

    doc.add_heading("Equation Support Test", level=1)
    doc.add_paragraph("Tests OMML → LaTeX conversion in MDM DOCX parser.")

    doc.add_heading("Block equation: Einstein", level=2)
    inject_math_block(doc, OMML_EINSTEIN)

    doc.add_heading("Inline equation: fraction", level=2)
    inject_math_inline(
        doc,
        "The probability is ",
        OMML_INLINE_FRAC,
        " when the coin is fair.",
    )

    doc.add_heading("Block equation: summation", level=2)
    inject_math_block(doc, OMML_SUM)

    doc.add_heading("Inline equation: cube root", level=2)
    inject_math_inline(
        doc,
        "Given ",
        OMML_SQRT,
        ", we compute.",
    )

    path = os.path.join(OUT_DIR, "test_equations.docx")
    doc.save(path)
    print(f"Created: {path}")
    return path


if __name__ == "__main__":
    create_equation_test()
