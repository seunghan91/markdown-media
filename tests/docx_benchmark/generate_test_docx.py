#!/usr/bin/env python3
"""Generate comprehensive test DOCX files for parser benchmarking."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))


def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph (python-docx doesn't natively support this)."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def create_comprehensive_test():
    """Test 1: Comprehensive document with all common features."""
    doc = Document()

    # --- Metadata ---
    doc.core_properties.title = "MDM DOCX Parser Benchmark"
    doc.core_properties.author = "Seunghan Kim"
    doc.core_properties.subject = "Parser Quality Test"

    # --- Heading hierarchy ---
    doc.add_heading("Document Title (H1)", level=1)
    doc.add_paragraph("This is body text under the main heading. It tests basic paragraph extraction.")

    doc.add_heading("Section One (H2)", level=2)
    doc.add_paragraph("Body text in section one.")

    doc.add_heading("Subsection 1.1 (H3)", level=3)
    doc.add_paragraph("Body text in subsection 1.1.")

    doc.add_heading("Sub-subsection 1.1.1 (H4)", level=4)
    doc.add_paragraph("Deep nested content.")

    # --- Inline formatting ---
    doc.add_heading("Formatting Tests (H2)", level=2)
    p = doc.add_paragraph()
    p.add_run("Bold text").bold = True
    p.add_run(", ")
    r = p.add_run("italic text")
    r.italic = True
    p.add_run(", ")
    r = p.add_run("bold italic")
    r.bold = True
    r.italic = True
    p.add_run(", ")
    r = p.add_run("strikethrough")
    r.font.strike = True
    p.add_run(", and normal text.")

    # Font size variation
    p2 = doc.add_paragraph()
    r = p2.add_run("Small text (8pt)")
    r.font.size = Pt(8)
    p2.add_run(" ")
    r = p2.add_run("Normal text (11pt)")
    r.font.size = Pt(11)
    p2.add_run(" ")
    r = p2.add_run("Large text (18pt)")
    r.font.size = Pt(18)

    # --- Bullet list ---
    doc.add_heading("List Tests (H2)", level=2)
    doc.add_paragraph("Bullet item 1", style="List Bullet")
    doc.add_paragraph("Bullet item 2", style="List Bullet")
    doc.add_paragraph("Bullet item 3", style="List Bullet")

    # --- Numbered list ---
    doc.add_paragraph("Numbered item 1", style="List Number")
    doc.add_paragraph("Numbered item 2", style="List Number")
    doc.add_paragraph("Numbered item 3", style="List Number")

    # --- Nested list (via indentation) ---
    p = doc.add_paragraph("Top level bullet", style="List Bullet")
    p2 = doc.add_paragraph("Nested bullet level 2", style="List Bullet 2")
    p3 = doc.add_paragraph("Nested bullet level 3", style="List Bullet 3")
    doc.add_paragraph("Back to top level", style="List Bullet")

    # --- Simple table ---
    doc.add_heading("Table Tests (H2)", level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ["Name", "Age", "City"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["Alice", "30", "Seoul"],
        ["Bob", "25", "Busan"],
        ["Charlie", "35", "Daejeon"],
    ]
    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            table.rows[r + 1].cells[c].text = val

    doc.add_paragraph("")  # spacer

    # --- Merged cell table ---
    doc.add_paragraph("Table with merged cells:")
    table2 = doc.add_table(rows=3, cols=3)
    table2.style = "Table Grid"
    # Horizontal merge: merge first row across 2 columns
    a = table2.rows[0].cells[0]
    b = table2.rows[0].cells[1]
    a.merge(b)
    a.text = "Merged A+B"
    table2.rows[0].cells[2].text = "C"
    # Vertical merge
    c1 = table2.rows[1].cells[0]
    c2 = table2.rows[2].cells[0]
    c1.merge(c2)
    c1.text = "Vertical Merge"
    table2.rows[1].cells[1].text = "D"
    table2.rows[1].cells[2].text = "E"
    table2.rows[2].cells[1].text = "F"
    table2.rows[2].cells[2].text = "G"

    # --- Hyperlinks ---
    doc.add_heading("Hyperlink Tests (H2)", level=2)
    p = doc.add_paragraph("Visit ")
    add_hyperlink(p, "https://github.com/seunghan91/markdown-media", "MDM GitHub Repository")
    p.add_run(" for more info.")

    p2 = doc.add_paragraph("Also check ")
    add_hyperlink(p2, "https://example.com", "Example Site")
    p2.add_run(".")

    # --- Footnotes (python-docx limited support, add via XML) ---
    doc.add_heading("Footnote Tests (H2)", level=2)
    p = doc.add_paragraph("This sentence has a footnote reference")
    # Add footnote reference via XML
    run = p.add_run()
    footnote_ref = OxmlElement("w:footnoteReference")
    footnote_ref.set(qn("w:id"), "1")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "FootnoteReference")
    rPr.append(rStyle)
    run._r.append(rPr)
    run._r.append(footnote_ref)

    # --- Korean content ---
    doc.add_heading("한국어 콘텐츠 테스트 (H2)", level=2)
    doc.add_paragraph("이것은 한국어 텍스트입니다. MDM 파서의 한국어 처리 능력을 테스트합니다.")

    p = doc.add_paragraph()
    r = p.add_run("굵은 한글")
    r.bold = True
    p.add_run(", ")
    r = p.add_run("기울임 한글")
    r.italic = True
    p.add_run(", ")
    r = p.add_run("굵은 기울임 한글")
    r.bold = True
    r.italic = True

    doc.add_paragraph("가) 첫 번째 항목")
    doc.add_paragraph("나) 두 번째 항목")
    doc.add_paragraph("다) 세 번째 항목")

    # --- Mixed Korean/English ---
    doc.add_heading("Mixed Language (H2)", level=2)
    p = doc.add_paragraph()
    p.add_run("The ")
    r = p.add_run("ActionCable")
    r.bold = True
    p.add_run(" 채널의 성능이 ")
    r = p.add_run("WebSocket")
    r.italic = True
    p.add_run(" 연결보다 느립니다.")

    # --- Blockquote (via style) ---
    doc.add_heading("Blockquote Test (H2)", level=2)
    try:
        doc.add_paragraph("This is a quote from an important source.", style="Quote")
    except:
        doc.add_paragraph("This is a quote from an important source.")

    # --- Horizontal rule approximation ---
    doc.add_paragraph("Content before rule.")
    doc.add_paragraph("---")  # Will test if parsers recognize this
    doc.add_paragraph("Content after rule.")

    path = os.path.join(OUT_DIR, "test_comprehensive.docx")
    doc.save(path)
    print(f"Created: {path}")
    return path


def create_table_stress_test():
    """Test 2: Complex tables that stress parsers."""
    doc = Document()
    doc.add_heading("Table Stress Test", level=1)

    # Simple 2x2
    doc.add_heading("Simple 2x2 Table", level=2)
    t = doc.add_table(rows=2, cols=2)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "Header A"
    t.rows[0].cells[1].text = "Header B"
    t.rows[1].cells[0].text = "Value 1"
    t.rows[1].cells[1].text = "Value 2"

    # Wide table (7 columns)
    doc.add_heading("Wide Table (7 columns)", level=2)
    t = doc.add_table(rows=3, cols=7)
    t.style = "Table Grid"
    for c in range(7):
        t.rows[0].cells[c].text = f"Col {c+1}"
    for r in range(1, 3):
        for c in range(7):
            t.rows[r].cells[c].text = f"R{r}C{c+1}"

    # Table with formatted content
    doc.add_heading("Table with Formatted Content", level=2)
    t = doc.add_table(rows=2, cols=2)
    t.style = "Table Grid"
    p = t.rows[0].cells[0].paragraphs[0]
    p.add_run("Bold header").bold = True
    t.rows[0].cells[1].text = "Normal header"
    p = t.rows[1].cells[0].paragraphs[0]
    r = p.add_run("Italic cell")
    r.italic = True
    t.rows[1].cells[1].text = "Cell with pipe | character"

    # Table with multiple paragraphs in a cell
    doc.add_heading("Table with Multi-paragraph Cells", level=2)
    t = doc.add_table(rows=2, cols=2)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    cell.text = "First paragraph"
    cell.add_paragraph("Second paragraph")
    t.rows[0].cells[1].text = "Single paragraph"
    t.rows[1].cells[0].text = "Simple cell"
    t.rows[1].cells[1].text = "Another cell"

    # Complex merge
    doc.add_heading("Complex Merged Table", level=2)
    t = doc.add_table(rows=4, cols=4)
    t.style = "Table Grid"
    # Header row - merge first two
    t.rows[0].cells[0].merge(t.rows[0].cells[1])
    t.rows[0].cells[0].text = "Group A"
    t.rows[0].cells[2].merge(t.rows[0].cells[3])
    t.rows[0].cells[2].text = "Group B"
    # Data rows
    for r in range(1, 4):
        for c in range(4):
            t.rows[r].cells[c].text = f"R{r}C{c+1}"
    # Vertical merge in column 0
    t.rows[1].cells[0].merge(t.rows[2].cells[0])
    t.rows[1].cells[0].text = "Vert Merged"

    path = os.path.join(OUT_DIR, "test_tables.docx")
    doc.save(path)
    print(f"Created: {path}")
    return path


def create_korean_government_style():
    """Test 3: Korean government document style."""
    doc = Document()

    doc.add_heading("대한민국 정부 공문서 (시험용)", level=1)

    doc.add_heading("제1조 (목적)", level=2)
    doc.add_paragraph(
        "이 규정은 MDM 파서의 한국어 공문서 변환 품질을 평가하기 위한 "
        "시험용 문서입니다."
    )

    doc.add_heading("제2조 (정의)", level=2)
    doc.add_paragraph(
        '이 규정에서 사용하는 용어의 뜻은 다음과 같다.'
    )
    doc.add_paragraph("1. \"파서\"란 문서를 구조적으로 분석하는 소프트웨어를 말한다.")
    doc.add_paragraph("2. \"마크다운\"이란 경량 마크업 언어를 말한다.")
    doc.add_paragraph("3. \"변환\"이란 한 형식에서 다른 형식으로 바꾸는 것을 말한다.")

    doc.add_heading("제3조 (적용 범위)", level=2)
    p = doc.add_paragraph("① 이 규정은 다음 각 호의 문서에 적용한다.")
    doc.add_paragraph("  1. HWP 형식의 공문서")
    doc.add_paragraph("  2. PDF 형식의 보고서")
    doc.add_paragraph("  3. DOCX 형식의 일반 문서")
    doc.add_paragraph("② 제1항에도 불구하고, 보안 문서는 이 규정의 적용 대상에서 제외한다.")

    # Korean table
    doc.add_heading("별표 1: 문서 형식 비교표", level=2)
    t = doc.add_table(rows=4, cols=4)
    t.style = "Table Grid"
    headers = ["형식", "개발사", "특징", "지원 여부"]
    for i, h in enumerate(headers):
        r = t.rows[0].cells[i].paragraphs[0].add_run(h)
        r.bold = True
    data = [
        ["HWP", "한컴", "한국 정부 표준", "✅"],
        ["PDF", "Adobe", "범용 문서", "✅"],
        ["DOCX", "Microsoft", "국제 표준", "✅"],
    ]
    for r_idx, row in enumerate(data):
        for c_idx, val in enumerate(row):
            t.rows[r_idx + 1].cells[c_idx].text = val

    doc.add_heading("부칙", level=2)
    doc.add_paragraph("이 규정은 2026년 4월 13일부터 시행한다.")

    path = os.path.join(OUT_DIR, "test_korean_gov.docx")
    doc.save(path)
    print(f"Created: {path}")
    return path


if __name__ == "__main__":
    create_comprehensive_test()
    create_table_stress_test()
    create_korean_government_style()
    print("\nAll test DOCX files created successfully!")
